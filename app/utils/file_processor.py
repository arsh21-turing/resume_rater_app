"""
Utility module for processing files in various formats.
Provides functions for reading, writing, and processing text files,
PDFs, Word documents, and other file formats.
"""

import os
import io
import PyPDF2
from io import BytesIO
import docx  # python-docx package

def read_file(file_object):
    """
    Read the content of a file or Streamlit's UploadedFile object based on its extension.
    
    Args:
        file_object: Can be either:
            - str: Path to a file on disk
            - UploadedFile: A Streamlit uploaded file object
            
    Returns:
        str: Content of the file as string
    """
    # Handle Streamlit's UploadedFile
    if hasattr(file_object, 'name') and hasattr(file_object, 'getvalue'):
        # This is a Streamlit UploadedFile
        file_name = file_object.name
        file_extension = os.path.splitext(file_name)[-1].lower()
        
        if file_extension == '.pdf':
            return read_pdf_stream(file_object)
        elif file_extension == '.docx':
            return read_docx_stream(file_object)
        elif file_extension == '.txt':
            return read_text_stream(file_object)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    # Handle file path string
    elif isinstance(file_object, str):
        file_path = file_object
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = os.path.splitext(file_path)[-1].lower()
        
        if file_extension == '.pdf':
            return read_pdf(file_path)
        elif file_extension == '.docx':
            return read_docx(file_path)
        elif file_extension == '.txt':
            return read_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    else:
        raise TypeError("Unsupported input type. Expected a file path string or Streamlit UploadedFile object.")

def read_text(file_path):
    """
    Read a plain text file.
    
    Args:
        file_path (str): Path to the text file
        
    Returns:
        str: Content of the text file
    """
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()
    
def export_results_to_csv(data, filename="resume_analysis_results.csv"):
    """
    Export resume analysis results to a CSV file.
    
    Args:
        data (dict): Analysis results data to export
        filename (str): Name of the output CSV file
    
    Returns:
        str: Path to the created CSV file or binary content as string
    """
    import csv
    import io
    
    # Create StringIO object to hold CSV data
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header row
    header = ["Category", "Score", "Details"]
    writer.writerow(header)
    
    # Write data rows
    if "skill_matches" in data:
        writer.writerow(["Skills Match", f"{data.get('skill_score', 0):.2f}", 
                         ", ".join(data.get("skill_matches", []))])
    
    if "experience" in data:
        writer.writerow(["Experience", f"{data.get('experience_score', 0):.2f}", 
                         f"{data.get('experience', {}).get('years', 0)} years"])
    
    if "education" in data:
        writer.writerow(["Education", f"{data.get('education_score', 0):.2f}", 
                         data.get("education", {}).get("highest_degree", "N/A")])
    
    # Write overall score
    writer.writerow(["Overall Match", f"{data.get('overall_score', 0):.2f}", 
                     f"{data.get('overall_score', 0) * 100:.1f}%"])
    
    # Get the CSV content as a string
    csv_content = output.getvalue()
    output.close()
    
    # Return the CSV content
    return csv_content

def read_pdf(file_path):
    """
    Extract text from a PDF file.
    
    Args:
        file_path (str): Path to the PDF file
        
    Returns:
        str: Extracted text from the PDF
    """
    text = ""
    with open(file_path, 'rb') as f:
        pdf_reader = PyPDF2.PdfReader(f)
        
        # Process each page
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text()
    
    return text

def read_docx(file_path):
    """
    Extract text from a DOCX file.
    
    Args:
        file_path (str): Path to the DOCX file
        
    Returns:
        str: Extracted text from the DOCX
    """
    doc = docx.Document(file_path)
    text = []
    
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    return '\n'.join(text)

# Additional functions for handling stream/memory file objects

def read_text_stream(file_object):
    """
    Read a plain text file from a stream.
    
    Args:
        file_object: A file-like object
        
    Returns:
        str: Content of the text file
    """
    if isinstance(file_object, io.TextIOBase):
        return file_object.read()
    else:
        return file_object.getvalue().decode('utf-8')

def read_pdf_stream(file_object):
    """
    Extract text from a PDF file stream.
    
    Args:
        file_object: A file-like object containing PDF data
        
    Returns:
        str: Extracted text from the PDF
    """
    text = ""
    pdf_reader = PyPDF2.PdfReader(file_object)
    
    # Process each page
    for page_num in range(len(pdf_reader.pages)):
        page = pdf_reader.pages[page_num]
        text += page.extract_text()
    
    return text

def read_docx_stream(file_object):
    """
    Extract text from a DOCX file stream.
    
    Args:
        file_object: A file-like object containing DOCX data
        
    Returns:
        str: Extracted text from the DOCX
    """
    doc = docx.Document(file_object)
    text = []
    
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    return '\n'.join(text)